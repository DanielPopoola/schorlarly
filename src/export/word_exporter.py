import logging
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from src.core.context_manager import ContextManager
from src.export.diagram_renderer import DiagramRenderer

logger = logging.getLogger(__name__)


CHAPTER_STRUCTURE = {
	'CHAPTER ONE: INTRODUCTION': [
		'Introduction',
		'Background to the Study',
		'Statement of the Problem',
		'Objective of the Study',
		'Significance of the Study',
		'Scope of the Study',
		'Limitations',
		'Organization of the Study',
		'Definition of Terms',
	],
	'CHAPTER TWO: LITERATURE REVIEW': [
		'Existing Approach to Problem Identified',
		'Effort to counter/solve existing challenges',
		'Specific Approach to Problem Identified',
	],
	'CHAPTER THREE: SYSTEM ANALYSIS AND DESIGN': [
		'System Analysis',
		'Method of Data Collection',
		'Problem of the Current System',
		'Objective of the new system',
		'Menu Specification',
		'Overview of the System Flowchart',
		'System Design',
		'Procedural Flowchart',
	],
	'CHAPTER FOUR: SYSTEM IMPLEMENTATION AND DOCUMENTATION': [
		'System Implementation',
		'System Requirement',
		'Hardware Requirement',
		'Software Requirement',
		'Test-Run',
		'Program Documentation',
		'User Manual',
		'System Maintenance',
	],
	'CHAPTER FIVE: SUMMARY, CONCLUSION AND RECOMMENDATION': ['Summary', 'Conclusion', 'Recommendations'],
}


def create_word_exporter_from_config(context_manager: ContextManager, config: dict) -> 'WordExporter':
	output_config = config.get('output', {})
	citation_config = config.get('citation', {})

	export_config = {'final_dir': output_config.get('final_dir', 'output/final'), 'citation': citation_config}

	return WordExporter(context_manager, export_config)


class WordExporter:
	def __init__(self, context_manager: ContextManager, config: dict):
		self.context_manager = context_manager
		self.output_dir = Path(config.get('final_dir', 'output/final'))
		self.citation_style = config.get('citation', {}).get('style', 'IEEE')

		diagrams_dir = self.output_dir / 'diagrams'
		self.diagram_renderer = DiagramRenderer(diagrams_dir)

		self.references_seen = set()

	def export(self, project_name: str, project_title: str, author: str) -> Path:
		logger.info(f'Exporting paper to Word: {project_name}')

		doc = Document()
		self._setup_document_style(doc)
		self._add_title_page(doc, project_title, author)
		section_contexts = [self.context_manager.get_section(name) for name in self.context_manager.section_order]
		rendered_diagrams = self.diagram_renderer.render_all_diagrams(section_contexts)
		self._add_chapters(doc, rendered_diagrams)
		self._add_references(doc)

		self.output_dir.mkdir(parents=True, exist_ok=True)
		output_path = self.output_dir / f'{project_name}.docx'
		doc.save(output_path)

		logger.info(f'Paper exported to: {output_path}')
		return output_path

	def _setup_document_style(self, doc: Document):
		style = doc.styles['Normal']
		font = style.font
		font.name = 'Times New Roman'
		font.size = Pt(12)

		paragraph_format = style.paragraph_format
		paragraph_format.line_spacing = 1.5
		paragraph_format.space_after = Pt(6)

		# Setup page margins (1 inch standard)
		sections = doc.sections
		for section in sections:
			section.top_margin = Inches(1)
			section.bottom_margin = Inches(1)
			section.left_margin = Inches(1)
			section.right_margin = Inches(1)

	def _add_chapters(self, doc, rendered_diagrams):
		for chapter_title, section_names in CHAPTER_STRUCTURE.items():
			doc.add_page_break()
			chapter_heading = doc.add_heading(chapter_title, 0)
			chapter_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

			for section_name in section_names:
				section_ctx = self.context_manager.get_section(section_name)
				if not section_ctx:
					logger.warning(f'Section not found: {section_name}')
					continue

				doc.add_heading(section_name, 1)
				cleaned_content = self._remove_embedded_references(section_ctx.content)
				self._add_markdown_content(doc, cleaned_content)

				if hasattr(section_ctx, 'diagrams') and section_ctx.diagrams:
					logger.debug(f'Section has {len(section_ctx.diagrams)} diagram(s)')
					for diagram in section_ctx.diagrams:
						diagram_key = f'{section_name}_{diagram["type"]}'
						if diagram_key in rendered_diagrams:
							self._add_diagram_to_doc(
								doc, rendered_diagrams[diagram_key]['path'], rendered_diagrams[diagram_key]['caption']
							)
						else:
							logger.warning(f'Diagram not rendered: {diagram_key}')

				doc.add_paragraph()

	def _add_diagram_to_doc(self, doc, image_path: Path, caption: str):
		if not image_path or not image_path.exists():
			logger.warning(f'Diagram not found: {image_path}')
			return

		from PIL import Image

		with Image.open(image_path) as img:
			width_px, height_px = img.size

		width_inches = width_px / 96
		height_inches = height_px / 96

		logger.info(f'Inserting diagram: {width_inches:.2f}" × {height_inches:.2f}"')

		doc.add_picture(str(image_path), width=Inches(6))

		caption_para = doc.add_paragraph(caption)
		caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
		caption_para.style = 'Caption'
		doc.add_paragraph()

	def _add_title_page(self, doc: Document, title: str, author: str):
		for _ in range(6):
			doc.add_paragraph()

		title_para = doc.add_heading(title, 0)
		title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

		doc.add_paragraph()  # Space

		# Author (centered)
		author_para = doc.add_paragraph(f'By\n{author}')
		author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

		doc.add_paragraph()

		date_para = doc.add_paragraph(datetime.now().strftime('%B %Y'))
		date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

		doc.add_page_break()

	def _remove_embedded_references(self, content: str) -> str:
		ref_header_pattern = r'(?:^|\n)#{1,3}\s*References?\s*\n.*'
		cleaned = re.sub(ref_header_pattern, '', content, flags=re.IGNORECASE | re.DOTALL)

		# Remove lines that look like references: [1] Author, "Title"...
		lines = cleaned.split('\n')
		cleaned_lines = []

		for line in lines:
			# Skip lines that start with [number] followed by author-like text
			if re.match(r'^\[\d+\]\s+[A-Z][\w\s,\.]+["\']', line.strip()):
				continue
			cleaned_lines.append(line)

		return '\n'.join(cleaned_lines)

	def _add_markdown_content(self, doc: Document, markdown_text: str):
		# Split into blocks by double newlines to maintain paragraph structure
		blocks = re.split(r'\n\n+', markdown_text)

		for block in blocks:
			block = block.strip()
			if not block:
				continue

			# 1. Handle Tables (process_tables adds to doc and returns empty string if matched)
			if '|' in block and '-' in block:
				remaining = self._process_tables(doc, block)
				if not remaining.strip():
					continue
				block = remaining.strip()

			# 2. Handle Code Blocks (detect_and_process_code_blocks adds to doc and returns empty string if matched)
			if block.startswith('```'):
				remaining = self._detect_and_process_code_blocks(doc, block)
				if not remaining.strip():
					continue
				block = remaining.strip()

			# 3. Handle Headers
			header_match = re.match(r'^(#{1,6})\s+(.+)$', block)
			if header_match:
				level = len(header_match.group(1))
				text = header_match.group(2).strip()
				doc.add_heading(text, level=level)
				continue

			# 4. Handle Lists
			if block.startswith('- ') or block.startswith('* '):
				self._add_bullet_list(doc, block)
			elif re.match(r'^\d+\.', block):
				self._add_numbered_list(doc, block)
			else:
				# 5. Regular Paragraph
				self._add_formatted_paragraph(doc, block)

	def _process_tables(self, doc: Document, text: str) -> str:
		table_pattern = r'\|(.+?)\|[\r\n]+\|[\s:|-]+\|[\r\n]+((?:\|.+?\|[\r\n]*)+)'

		def replace_table(match):
			try:
				header_line = match.group(1)
				body_lines = match.group(2)

				headers = [cell.strip() for cell in header_line.split('|') if cell.strip()]

				if not headers:
					return match.group(0)

				rows = []
				for line in body_lines.strip().split('\n'):
					if not line.strip() or line.strip().startswith('|---'):
						continue

					cells = [cell.strip() for cell in line.split('|') if cell.strip()]
					if cells and len(cells) == len(headers):
						rows.append(cells)

				if not rows:
					return match.group(0)

				table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
				table.style = 'Light Grid Accent 1'

				for i, header in enumerate(headers):
					cell = table.rows[0].cells[i]
					cell.text = header
					for paragraph in cell.paragraphs:
						for run in paragraph.runs:
							run.bold = True

				for row_idx, row_data in enumerate(rows, start=1):
					for col_idx, cell_data in enumerate(row_data):
						table.rows[row_idx].cells[col_idx].text = cell_data

				doc.add_paragraph()
				return ''

			except Exception as e:
				logger.error(f'Table conversion failed: {e}')
				return match.group(0)  # Return original on error

		text = re.sub(table_pattern, replace_table, text, flags=re.MULTILINE)
		return text

	def _add_formatted_paragraph(self, doc: Document, text: str):
		para = doc.add_paragraph()

		pattern = r'(\*\*\*[^*]+\*\*\*|\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)'
		parts = re.split(pattern, text)

		for part in parts:
			if not part:
				continue

			if part.startswith('***') and part.endswith('***'):
				run = para.add_run(part[3:-3])
				run.bold = True
				run.italic = True

			elif part.startswith('**') and part.endswith('**'):
				run = para.add_run(part[2:-2])
				run.bold = True

			elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
				run = para.add_run(part[1:-1])
				run.italic = True

			elif part.startswith('`') and part.endswith('`'):
				run = para.add_run(part[1:-1])
				run.font.name = 'Courier New'
				run.font.size = Pt(10)

			else:
				para.add_run(part)

	def _add_bullet_list(self, doc: Document, text: str):
		lines = text.split('\n')
		for line in lines:
			line = line.strip()
			if line.startswith('- ') or line.startswith('* '):
				content = line[2:]
				doc.add_paragraph(content, style='List Bullet')

	def _add_numbered_list(self, doc: Document, text: str):
		lines = text.split('\n')
		for line in lines:
			line = line.strip()
			if re.match(r'^\d+\.', line):
				content = re.sub(r'^\d+\.\s*', '', line)
				doc.add_paragraph(content, style='List Number')

	def _detect_and_process_code_blocks(self, doc: Document, text: str) -> str:
		"""Extract code blocks and replace with Word elements"""

		# Pattern that handles variations
		code_block_pattern = r'```(?P<lang>\w*)\s*\n(?P<code>.*?)\n\s*```'

		def replace_code(match):
			lang = match.group('lang') or 'code'
			code = match.group('code')

			# Add to document
			para = doc.add_paragraph(code.strip())
			para.style = 'Normal'

			for run in para.runs:
				run.font.name = 'Courier New'
				run.font.size = Pt(10)

			# Add language label if provided
			if lang != 'code':
				label_para = doc.add_paragraph(f'({lang})', style='Caption')
				label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

			from docx.oxml import OxmlElement
			from docx.oxml.ns import qn

			shading_elm = OxmlElement('w:shd')
			shading_elm.set(qn('w:fill'), 'F0F0F0')
			para._element.get_or_add_pPr().append(shading_elm)

			return ''  # Remove from text

		text = re.sub(code_block_pattern, replace_code, text, flags=re.DOTALL)
		return text

	def _add_references(self, doc: Document):
		if not self.context_manager.citation_registry:
			logger.info('No citations in registry')
			return

		doc.add_page_break()
		ref_heading = doc.add_heading('REFERENCES', 0)
		ref_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

		# Sort by citation number
		sorted_citations = sorted(
			self.context_manager.citation_registry.items(),
			key=lambda x: x[1][0],  # Sort by number (second element of tuple)
		)

		for _, (number, paper_dict) in sorted_citations:
			reference_text = self._format_reference_ieee(number, paper_dict)
			doc.add_paragraph(reference_text, style='Normal')

		logger.info(f'Added {len(sorted_citations)} references')

	def _format_reference_ieee(self, number: int, paper_dict: dict) -> str:
		"""Format a single reference in IEEE style."""
		authors = paper_dict.get('authors', [])

		# Format authors: J. Smith, A. Doe
		if authors:
			author_str = ', '.join(
				[
					f'{a.split()[-1]}, {a.split()[0][0]}.' if ' ' in a else a
					for a in authors[:3]  # Max 3 authors
				]
			)
			if len(authors) > 3:
				author_str += ', et al.'
		else:
			author_str = 'Unknown'

		title = paper_dict.get('title', 'Untitled')
		year = paper_dict.get('year', 'n.d.')
		url = paper_dict.get('url', '')

		# IEEE format: [1] A. Smith, "Title," Year. [Online]. Available: URL
		reference = f'[{number}] {author_str}, "{title}," {year}.'

		if url:
			reference += f' [Online]. Available: {url}'

		return reference

	def _extract_full_references(self) -> dict[str, str]:
		references = {}
		all_citations = set()

		# Step 1: Collect all citation markers from all sections
		for section_name in self.context_manager.section_order:
			section_ctx = self.context_manager.get_section(section_name)
			if not section_ctx:
				continue

			# Find all [N] patterns
			citations = re.findall(r'\[(\d+)\]', section_ctx.content)
			all_citations.update(citations)

		# Step 2: Try to find full reference text for each citation
		for section_name in self.context_manager.section_order:
			section_ctx = self.context_manager.get_section(section_name)
			if not section_ctx:
				continue

			content = section_ctx.content

			# Look for lines that start with [N] followed by reference text
			# Example: [1] Author, "Title", Journal, Year
			pattern = r'^\[(\d+)\]\s+(.+?)(?=^\[\d+\]|\Z)'
			matches = re.finditer(pattern, content, re.MULTILINE | re.DOTALL)

			for match in matches:
				num = match.group(1)
				ref_text = match.group(2).strip()

				ref_text = re.sub(r'\s+', ' ', ref_text)

				if len(ref_text) > 20:
					references[num] = f'[{num}] {ref_text}'

		if all_citations and not references:
			logger.warning('Citations found but no full references - creating placeholders')
			for num in sorted(all_citations, key=int):
				references[num] = f'[{num}] Reference details not available'

		return references
